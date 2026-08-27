import io
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.enums import TA_RIGHT, TA_LEFT, TA_CENTER
from reportlab.lib import colors

# pyrefly: ignore [missing-import]
from docx import Document
# pyrefly: ignore [missing-import]
from docx.shared import Inches, Pt
# pyrefly: ignore [missing-import]
from docx.enum.text import WD_ALIGN_PARAGRAPH
# pyrefly: ignore [missing-import]
from docx.oxml import OxmlElement
# pyrefly: ignore [missing-import]
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# PDF Generation using ReportLab (Pure Python, Cross-Platform)
# ---------------------------------------------------------------------------

def _draw_page_number_factory(has_title_page: bool):
    """
    Returns a canvas callback for drawing standard screenplay header page numbers.
    """
    def _callback(canvas, doc):
        # Calculate actual screenplay script page
        script_page = (doc.page - 1) if has_title_page else doc.page
        if script_page > 1:
            canvas.saveState()
            canvas.setFont("Courier", 12)
            # 1.0 inch from right edge, 0.75 inch from top edge
            x = 8.5 * inch - 1.0 * inch
            y = 11.0 * inch - 0.75 * inch
            canvas.drawRightString(x, y, f"{script_page}.")
            canvas.restoreState()

    return _callback


def export_script_to_pdf(script) -> bytes:
    """
    Renders script lines and optional Title Page into a paginated screenplay-format PDF byte stream.
    Supports Title Page cover formatting and side-by-side Dual Dialogue rendering.
    """
    buffer = io.BytesIO()
    
    # Standard Screenplay Document Template (Letter, 1.5" Left Margin, 1.0" Right/Top/Bottom)
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=1.5 * inch,
        rightMargin=1.0 * inch,
        topMargin=1.0 * inch,
        bottomMargin=1.0 * inch,
    )
    
    styles = getSampleStyleSheet()
    
    # Custom Screenplay Paragraph Styles (Courier 12pt, 14pt leading)
    style_heading = ParagraphStyle(
        'ScreenplaySceneHeading',
        fontName='Courier-Bold',
        fontSize=12,
        leading=14,
        leftIndent=0,
        rightIndent=0,
        spaceBefore=24,
        spaceAfter=0,
        keepWithNext=True,
    )
    
    style_action = ParagraphStyle(
        'ScreenplayAction',
        fontName='Courier',
        fontSize=12,
        leading=14,
        leftIndent=0,
        rightIndent=0,
        spaceBefore=12,
        spaceAfter=0,
    )
    
    style_character = ParagraphStyle(
        'ScreenplayCharacter',
        fontName='Courier-Bold',
        fontSize=12,
        leading=14,
        leftIndent=2.2 * inch,
        rightIndent=0,
        spaceBefore=12,
        spaceAfter=0,
        keepWithNext=True,
    )
    
    style_parenthetical = ParagraphStyle(
        'ScreenplayParenthetical',
        fontName='Courier',
        fontSize=12,
        leading=14,
        leftIndent=1.6 * inch,
        rightIndent=1.9 * inch,
        spaceBefore=0,
        spaceAfter=0,
        keepWithNext=True,
    )
    
    style_dialogue = ParagraphStyle(
        'ScreenplayDialogue',
        fontName='Courier',
        fontSize=12,
        leading=14,
        leftIndent=1.0 * inch,
        rightIndent=1.5 * inch,
        spaceBefore=0,
        spaceAfter=6,
    )
    
    style_transition = ParagraphStyle(
        'ScreenplayTransition',
        fontName='Courier-Bold',
        fontSize=12,
        leading=14,
        alignment=TA_RIGHT,
        spaceBefore=12,
        spaceAfter=12,
    )

    # Title Page Styles
    style_tp_title = ParagraphStyle(
        'TPTitle',
        fontName='Courier-Bold',
        fontSize=18,
        leading=22,
        alignment=TA_CENTER,
        spaceBefore=0,
        spaceAfter=12,
    )
    style_tp_credit = ParagraphStyle(
        'TPCredit',
        fontName='Courier',
        fontSize=12,
        leading=16,
        alignment=TA_CENTER,
        spaceBefore=12,
        spaceAfter=6,
    )
    style_tp_author = ParagraphStyle(
        'TPAuthor',
        fontName='Courier',
        fontSize=13,
        leading=16,
        alignment=TA_CENTER,
        spaceBefore=6,
        spaceAfter=12,
    )
    style_tp_source = ParagraphStyle(
        'TPSource',
        fontName='Courier-Oblique',
        fontSize=11,
        leading=14,
        alignment=TA_CENTER,
        spaceBefore=6,
        spaceAfter=12,
    )
    style_tp_bottom = ParagraphStyle(
        'TPBottom',
        fontName='Courier',
        fontSize=10,
        leading=13,
        alignment=TA_LEFT,
        spaceBefore=4,
        spaceAfter=0,
    )

    # Dual Dialogue Table Styles
    style_dual_char = ParagraphStyle(
        'DualChar',
        fontName='Courier-Bold',
        fontSize=11,
        leading=13,
        alignment=TA_CENTER,
        spaceBefore=6,
        spaceAfter=0,
    )
    style_dual_paren = ParagraphStyle(
        'DualParen',
        fontName='Courier',
        fontSize=10,
        leading=12,
        alignment=TA_CENTER,
        spaceBefore=0,
        spaceAfter=0,
    )
    style_dual_dial = ParagraphStyle(
        'DualDial',
        fontName='Courier',
        fontSize=11,
        leading=13,
        alignment=TA_LEFT,
        spaceBefore=0,
        spaceAfter=4,
    )
    
    style_map = {
        "scene_heading": style_heading,
        "action": style_action,
        "character": style_character,
        "parenthetical": style_parenthetical,
        "dialogue": style_dialogue,
        "transition": style_transition,
    }

    story = []
    has_title_page = False

    # 1. Build Title Page if metadata exists
    title_page = getattr(script, "title_page", None)
    if title_page and (title_page.title or title_page.author or title_page.contact):
        has_title_page = True
        story.append(Spacer(1, 2.5 * inch))
        title_text = (title_page.title or script.title).upper()
        story.append(Paragraph(title_text, style_tp_title))
        
        credit_text = title_page.credit or "written by"
        story.append(Paragraph(credit_text, style_tp_credit))

        author_text = title_page.author or "Anonymous"
        story.append(Paragraph(author_text, style_tp_author))

        if title_page.source:
            story.append(Paragraph(title_page.source, style_tp_source))

        story.append(Spacer(1, 3.2 * inch))

        if title_page.draft_date:
            story.append(Paragraph(f"Draft Date: {title_page.draft_date}", style_tp_bottom))
        if title_page.contact:
            for line in title_page.contact.splitlines():
                if line.strip():
                    story.append(Paragraph(line.strip(), style_tp_bottom))
        if title_page.copyright:
            story.append(Paragraph(title_page.copyright, style_tp_bottom))

        story.append(PageBreak())

    # 2. Build Screenplay Scenes and Dual Dialogue
    scenes = sorted(script.scenes.prefetch_related("lines").all(), key=lambda s: s.order)
    is_first = True
    
    for scene in scenes:
        lines = sorted(scene.lines.all(), key=lambda l: l.order)
        idx = 0

        while idx < len(lines):
            line = lines[idx]

            # Check for Dual Dialogue block (starting with left)
            if line.is_dual_dialogue and line.dual_pos == "left":
                left_flowables = []
                right_flowables = []

                # Collect left speaker lines
                while idx < len(lines) and lines[idx].is_dual_dialogue and lines[idx].dual_pos == "left":
                    l_item = lines[idx]
                    esc_text = l_item.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    if l_item.type == "character":
                        left_flowables.append(Paragraph(esc_text.upper(), style_dual_char))
                    elif l_item.type == "parenthetical":
                        left_flowables.append(Paragraph(esc_text, style_dual_paren))
                    else:
                        left_flowables.append(Paragraph(esc_text, style_dual_dial))
                    idx += 1

                # Collect right speaker lines
                while idx < len(lines) and lines[idx].is_dual_dialogue and lines[idx].dual_pos == "right":
                    r_item = lines[idx]
                    esc_text = r_item.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    if r_item.type == "character":
                        right_flowables.append(Paragraph(esc_text.upper(), style_dual_char))
                    elif r_item.type == "parenthetical":
                        right_flowables.append(Paragraph(esc_text, style_dual_paren))
                    else:
                        right_flowables.append(Paragraph(esc_text, style_dual_dial))
                    idx += 1

                # Render side-by-side 2-column Table
                dual_table = Table(
                    [[left_flowables, right_flowables]],
                    colWidths=[2.85 * inch, 2.85 * inch],
                )
                dual_table.setStyle(
                    TableStyle([
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 0),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                        ('TOPPADDING', (0, 0), (-1, -1), 4),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ])
                )
                story.append(dual_table)
                is_first = False
                continue

            # Standard line
            st = style_map.get(line.type, style_action)
            raw_text = line.text

            # Scene heading with scene number
            if line.type == "scene_heading" and scene.scene_number and f"#{scene.scene_number}#" not in raw_text:
                raw_text = f"{raw_text} #{scene.scene_number}#"

            # Escape HTML characters for ReportLab XML rendering
            text_escaped = (
                raw_text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            
            if is_first and not has_title_page:
                # Remove space before top line on page 1
                first_style = ParagraphStyle('FirstStyle', parent=st, spaceBefore=0)
                story.append(Paragraph(text_escaped, first_style))
                is_first = False
            else:
                story.append(Paragraph(text_escaped, st))
                is_first = False

            idx += 1

    # Build PDF with page numbers on later pages
    page_callback = _draw_page_number_factory(has_title_page)
    doc.build(story, onFirstPage=lambda c, d: None, onLaterPages=page_callback)
    
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes


# ---------------------------------------------------------------------------
# Word Document (.docx) Generation using python-docx
# ---------------------------------------------------------------------------

def add_page_number(run):
    """
    Appends a dynamic PAGE field code to a Word run.
    """
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)


def export_script_to_word(script) -> bytes:
    """
    Renders script lines and optional Title Page into a formatted Word (.docx) file.
    Supports Title Page cover formatting and side-by-side Dual Dialogue tables.
    """
    doc = Document()
    
    # Page setup: Standard Letter size
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    
    # Standard screenplay margins (1.5" left for binding, 1.0" top/bottom/right)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)
    
    # Header page numbers in top right corner (starting from page 2)
    section.different_first_page_header_footer = True
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_p.add_run()
    header_run.font.name = 'Courier New'
    header_run.font.size = Pt(12)
    add_page_number(header_run)
    header_p.add_run(".")

    has_title_page = False
    title_page = getattr(script, "title_page", None)

    # 1. Render Title Page if metadata exists
    if title_page and (title_page.title or title_page.author or title_page.contact):
        has_title_page = True
        
        # Title (centered, bold, 18pt)
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(180)
        p_title.paragraph_format.space_after = Pt(12)
        r_title = p_title.add_run((title_page.title or script.title).upper())
        r_title.font.name = 'Courier New'
        r_title.font.size = Pt(18)
        r_title.font.bold = True

        # Credit
        p_credit = doc.add_paragraph()
        p_credit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_credit.paragraph_format.space_before = Pt(12)
        p_credit.paragraph_format.space_after = Pt(6)
        r_credit = p_credit.add_run(title_page.credit or "written by")
        r_credit.font.name = 'Courier New'
        r_credit.font.size = Pt(12)

        # Author
        p_author = doc.add_paragraph()
        p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_author.paragraph_format.space_before = Pt(6)
        p_author.paragraph_format.space_after = Pt(12)
        r_author = p_author.add_run(title_page.author or "Anonymous")
        r_author.font.name = 'Courier New'
        r_author.font.size = Pt(13)

        if title_page.source:
            p_source = doc.add_paragraph()
            p_source.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_source.paragraph_format.space_before = Pt(6)
            p_source.paragraph_format.space_after = Pt(12)
            r_source = p_source.add_run(title_page.source)
            r_source.font.name = 'Courier New'
            r_source.font.size = Pt(11)
            r_source.font.italic = True

        # Bottom metadata block
        p_bottom = doc.add_paragraph()
        p_bottom.paragraph_format.space_before = Pt(200)
        p_bottom.paragraph_format.space_after = Pt(0)
        p_bottom.paragraph_format.line_spacing = 1.1

        bottom_lines = []
        if title_page.draft_date:
            bottom_lines.append(f"Draft Date: {title_page.draft_date}")
        if title_page.contact:
            bottom_lines.extend(title_page.contact.splitlines())
        if title_page.copyright:
            bottom_lines.append(title_page.copyright)

        for bl in bottom_lines:
            if bl.strip():
                p_b = doc.add_paragraph()
                p_b.paragraph_format.space_before = Pt(2)
                p_b.paragraph_format.space_after = Pt(0)
                r_b = p_b.add_run(bl.strip())
                r_b.font.name = 'Courier New'
                r_b.font.size = Pt(10)

        doc.add_page_break()
    
    def add_paragraph_element(text, element_type):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        
        # Setup indents relative to section margins (left=1.5", right=1.0")
        if element_type == "scene_heading":
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.right_indent = Inches(0)
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_with_next = True
        elif element_type == "action":
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.right_indent = Inches(0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(0)
        elif element_type == "character":
            p.paragraph_format.left_indent = Inches(2.2)
            p.paragraph_format.right_indent = Inches(0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_with_next = True
        elif element_type == "parenthetical":
            p.paragraph_format.left_indent = Inches(1.6)
            p.paragraph_format.right_indent = Inches(1.9)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_with_next = True
        elif element_type == "dialogue":
            p.paragraph_format.left_indent = Inches(1.0)
            p.paragraph_format.right_indent = Inches(1.5)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
        elif element_type == "transition":
            p.paragraph_format.left_indent = Inches(4.0)
            p.paragraph_format.right_indent = Inches(0)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
        else:
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.space_before = Pt(6)
            
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(12)
        
        # Bold styling for scene headings/characters/transitions
        if element_type in ("scene_heading", "character", "transition"):
            run.font.bold = True
            
        return p

    # Pull scenes and lines
    scenes = sorted(script.scenes.prefetch_related("lines").all(), key=lambda s: s.order)
    is_first = True
    
    for scene in scenes:
        lines = sorted(scene.lines.all(), key=lambda l: l.order)
        idx = 0

        while idx < len(lines):
            line = lines[idx]

            # Check for Dual Dialogue block in Word
            if line.is_dual_dialogue and line.dual_pos == "left":
                left_lines = []
                right_lines = []

                while idx < len(lines) and lines[idx].is_dual_dialogue and lines[idx].dual_pos == "left":
                    left_lines.append(lines[idx])
                    idx += 1

                while idx < len(lines) and lines[idx].is_dual_dialogue and lines[idx].dual_pos == "right":
                    right_lines.append(lines[idx])
                    idx += 1

                # Create 1-row, 2-column Table
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                table.columns[0].width = Inches(2.9)
                table.columns[1].width = Inches(2.9)

                # Populate Left Column
                cell_left = table.cell(0, 0)
                cell_left.width = Inches(2.9)
                p_l = cell_left.paragraphs[0]
                for l_idx, l_line in enumerate(left_lines):
                    if l_idx > 0:
                        p_l = cell_left.add_paragraph()
                    p_l.paragraph_format.space_before = Pt(0)
                    p_l.paragraph_format.space_after = Pt(2)
                    if l_line.type == "character":
                        p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p_l.add_run(l_line.text.upper())
                        r.font.bold = True
                    elif l_line.type == "parenthetical":
                        p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p_l.add_run(l_line.text)
                    else:
                        p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        r = p_l.add_run(l_line.text)
                    r.font.name = 'Courier New'
                    r.font.size = Pt(10)

                # Populate Right Column
                cell_right = table.cell(0, 1)
                cell_right.width = Inches(2.9)
                p_r = cell_right.paragraphs[0]
                for r_idx, r_line in enumerate(right_lines):
                    if r_idx > 0:
                        p_r = cell_right.add_paragraph()
                    p_r.paragraph_format.space_before = Pt(0)
                    p_r.paragraph_format.space_after = Pt(2)
                    if r_line.type == "character":
                        p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p_r.add_run(r_line.text.upper())
                        r.font.bold = True
                    elif r_line.type == "parenthetical":
                        p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p_r.add_run(r_line.text)
                    else:
                        p_r.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        r = p_r.add_run(r_line.text)
                    r.font.name = 'Courier New'
                    r.font.size = Pt(10)

                is_first = False
                continue

            text_to_render = line.text
            if line.type == "scene_heading" and scene.scene_number and f"#{scene.scene_number}#" not in text_to_render:
                text_to_render = f"{text_to_render} #{scene.scene_number}#"

            p = add_paragraph_element(text_to_render, line.type)
            if is_first and not has_title_page:
                p.paragraph_format.space_before = Pt(0)
                is_first = False
            else:
                is_first = False

            idx += 1
                
    # Save document to memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()
